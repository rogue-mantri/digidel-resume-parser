import os
import re
import json
import csv
from datetime import datetime
from pathlib import Path

class ResumeExtractor:
    SUPPORTED_FORMATS = {'.pdf', '.docx', '.doc', '.txt', '.rtf'}
    def __init__(self):
        self.stats = {'processed': 0, 'failed': 0, 'by_format': {}}
    def extract(self, file_path):
        file_path = Path(file_path)
        ext = file_path.suffix.lower()
        if ext not in self.SUPPORTED_FORMATS:
            return {'success': False, 'error': f'Unsupported format: {ext}', 'text': ''}
        try:
            if ext == '.pdf': text = self._extract_pdf(file_path)
            elif ext in {'.docx', '.doc'}: text = self._extract_docx(file_path)
            elif ext == '.txt': text = self._extract_txt(file_path)
            elif ext == '.rtf': text = self._extract_rtf(file_path)
            else: text = ''
            self.stats['processed'] += 1
            self.stats['by_format'][ext] = self.stats['by_format'].get(ext, 0) + 1
            return {'success': True, 'text': self._clean_text(text), 'raw_text': text, 'file_name': file_path.name, 'file_size': file_path.stat().st_size, 'format': ext}
        except Exception as e:
            self.stats['failed'] += 1
            return {'success': False, 'error': str(e), 'text': ''}
    def _extract_pdf(self, file_path):
        try:
            from pypdf import PdfReader
            reader = PdfReader(str(file_path))
            return "\n".join(page.extract_text() or "" for page in reader.pages)
        except ImportError:
            try:
                import PyPDF2
                with open(file_path, 'rb') as f:
                    reader = PyPDF2.PdfReader(f)
                    return "\n".join(page.extract_text() or "" for page in reader.pages)
            except ImportError:
                raise ImportError("Neither pypdf nor PyPDF2 is installed. Run: pip install pypdf")
    def _extract_docx(self, file_path):
        try:
            from docx import Document
            doc = Document(str(file_path))
            text = "\n".join(paragraph.text for paragraph in doc.paragraphs)
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        text += "\n" + cell.text
            return text
        except ImportError:
            raise ImportError("python-docx is not installed. Run: pip install python-docx")
    def _extract_txt(self, file_path):
        encodings = ['utf-8', 'latin-1', 'cp1252', 'iso-8859-1']
        for encoding in encodings:
            try:
                with open(file_path, 'r', encoding=encoding) as f:
                    return f.read()
            except UnicodeDecodeError:
                continue
        raise UnicodeDecodeError(f"Could not decode file with any encoding: {file_path}")
    def _extract_rtf(self, file_path):
        with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
            content = f.read()
        text = re.sub(r'\\[a-z]+\d*\s?', ' ', content)
        text = re.sub(r'[{}]', '', text)
        text = re.sub(r'\\[\\\{\}\|]', '', text)
        return text
    def _clean_text(self, text):
        text = text.replace('â\x80\x94', '--').replace('â\x80\x93', '-')
        text = text.replace('â\x80\x9c', '"').replace('â\x80\x9d', '"')
        text = text.replace('â\x80\x99', "'").replace('â\x80\xa2', '*')
        text = re.sub(r'[ \t]+', ' ', text)
        text = '\n'.join(line.strip() for line in text.split('\n'))
        return text
    def get_stats(self):
        return self.stats

def extract_batch(folder_path, output_dir=None):
    folder = Path(folder_path)
    extractor = ResumeExtractor()
    results = []
    for file_path in folder.iterdir():
        if file_path.suffix.lower() in ResumeExtractor.SUPPORTED_FORMATS:
            result = extractor.extract(file_path)
            result['file_path'] = str(file_path)
            results.append(result)
    if output_dir:
        output_dir = Path(output_dir)
        output_dir.mkdir(parents=True, exist_ok=True)
        for result in results:
            if result['success']:
                out_file = output_dir / f"{result['file_name']}.txt"
                with open(out_file, 'w', encoding='utf-8') as f:
                    f.write(result['text'])
    return results, extractor.get_stats()
